# -*- coding: utf-8 -*-
"""
Created on Thu Jul 23 22:30:21 2020

@author: pc
"""
import streamlit as st
import re
from string import punctuation
from nltk import word_tokenize
from nltk.corpus import stopwords
import docx
import time

def findclosestsentence(AllParagraphs,ParagraphNo,wordsFiltered2):
    closestsentence:str=" "
    maxwordnum:int=0
    i:int=0
    wordsofclosestsentence=list()
    txt=AllParagraphs[ParagraphNo-1]
    sentences=re.split(r' *[\.\?!][\'"\)\]]* *', txt)
    for stuff in sentences:
        matchingwords:int=0
        kelime = re.split(r'\W+', stuff)
        listToStr = ' '.join([str(elem.lower()) for elem in kelime])
        tokens = word_tokenize(listToStr)
        stoplist = stopwords.words('english') + list(punctuation)
        wordsFiltered3 = [token for token in tokens if token not in stoplist]
        for word in wordsFiltered2:
           if word in wordsFiltered3:
               matchingwords+=1
        if matchingwords>maxwordnum:
            maxwordnum=matchingwords
            closestsentence=stuff.lower()
            wordsofclosestsentence=wordsFiltered3
        st.write("Number of words matching sentence "+str(i+1)+": "+str(matchingwords))
        i+=1
    
    if closestsentence==" ":
      st.error("There is no answer for this question")
    else:
      st.warning("The closest sentence to this question: "+closestsentence)    
    return wordsofclosestsentence         

def findwordsinarticle(AllParagraphs):
    WordsofArticle=list()
    UniqueWords=list()
    for a in range(0,len(AllParagraphs)-1):
        txt=AllParagraphs[a]
        sentences=re.split(r' *[\.\?!][\'"\)\]]* *', txt)
        for stuff in sentences:
           kelime = re.split(r'\W+', stuff)
           listToStr = ' '.join([str(elem) for elem in kelime])
           tokens = word_tokenize(listToStr)
           stoplist = stopwords.words('english') + list(punctuation)
           wordsFiltered1 = [token for token in tokens if token not in stoplist] 
           for word in sorted(wordsFiltered1):
               if word not in UniqueWords:
                  UniqueWords.append(word)
           WordsofArticle.append(UniqueWords)
    return WordsofArticle

def findtheanswerofquestions(ParagraphNo,QuestionNo,data2):
    AnswersDividedByParagraphs=data2.split("%")
    AnswersofQuestions=AnswersDividedByParagraphs[ParagraphNo-1].split("*")
    return AnswersofQuestions[QuestionNo-1]

def findthequestionofparagraphs(ParagraphNo,QuestionNo,data):#paragrafın sorularından indeksle çağırdığım soru dönecek
    
    QuestionsDivideByParagraphs=data.split("$")
    QuestionsofParagraphs=QuestionsDivideByParagraphs[ParagraphNo-1].split("?") 
    return QuestionsofParagraphs[QuestionNo-1]
    
def extractquestion(Question,QuestionNo):
    kelime = re.split(r'\W+', Question)    
    listToStr = ' '.join([str(elem) for elem in kelime])
    tokens = word_tokenize(listToStr)
    stoplist = stopwords.words('english') + list(punctuation)
    wordsFiltered2 = [token for token in tokens if token not in stoplist]
    st.write("Words of question "+str(QuestionNo))
    for word in sorted(wordsFiltered2):
        st.write(word)
    return wordsFiltered2

def choicefile(choice):
    dosya=choice+".docx"
    return dosya
def choicequestion(choice):
    question=choice+"Questions.txt"
    return question
def choiceanswer(choice):
    answer=choice+"Answers.txt"
    return answer

st.title("Streamlit NLP App")
st.balloons()
article= ["Choose one ","Immune System","Steam Engine"]
st.sidebar.header("Article")
choice = st.sidebar.selectbox(" ",article)
choice = choice.replace(' ', '')
dosya:str=" "
question:str=" "
answer:str=" "#bir tane metot yapılıp articledan seçilen isim yollanıp choice bu eşitlemeler metot içinde yapılabilir
dosya=choicefile(choice)
question=choicequestion(choice)
answer=choiceanswer(choice)    
f = open(question, "r")#sorular dosyasını açar
data=f.read().lower()#veriyi dataya eşitler
d=open(answer,"r")
data2=d.read().lower()
doc =docx.Document(dosya)
all_paras=doc.paragraphs
txt:str=" "
AllParagraphs=list()#articleın paragraflara bölünmüş hali
remainingwords=list()
for para in range(0,len(all_paras)-1):
   AllParagraphs.append(doc.paragraphs[para].text)
st.sidebar.subheader('Select a value for paragraph number')  
#ParagraphNo = st.sidebar.slider(" ",1,50)
ParagraphNo = st.sidebar.number_input("Paragraph Number", min_value=0,value=50)
st.sidebar.subheader('Select a value for question number')
QuestionNo = st.sidebar.number_input("Question Number", min_value=0,value=50)
#QuestionNo=st.sidebar.slider("  ",1,50)

if st.sidebar.button("Show"):
  paragraph=doc.paragraphs[ParagraphNo-1].text
  Question=findthequestionofparagraphs(ParagraphNo,QuestionNo,data)
  Answer=findtheanswerofquestions(ParagraphNo, QuestionNo, data2)
  st.success(paragraph)
  st.warning("Question: "+Question+"?")
  st.warning("Answer: "+Answer)
  wordsFiltered2=extractquestion(Question, QuestionNo)#sorunun kelimeleri
  wordsofclosestsentence=findclosestsentence(AllParagraphs, ParagraphNo, wordsFiltered2)
  with st.spinner("Checking if the answer is in the closest sentence or not..."):
      time.sleep(5)
  st.info("Done!")
  st.write("Remains after removing words from the closest sentence")
  for w in wordsofclosestsentence:
     if w not in wordsFiltered2:
         if w not in remainingwords:
           remainingwords.append(w)
  for p in remainingwords:
      if p==Answer:
          st.success(Answer)
      if p!=Answer:
          st.write(p)
  if Answer in remainingwords:
     st.success("Answer is in the closest sentence!")
  else:
      st.error("The closest sentence does not have answer of this question!")


